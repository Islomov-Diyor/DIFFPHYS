import os
import json
import random
import difflib
import requests

from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage

from .forms import PronunciationForm, TextAnalysisForm

from docx import Document
import fitz  # PyMuPDF
from decouple import config

import speech_recognition as sr
from pydub import AudioSegment
from gtts import gTTS
from django.http import FileResponse


# =============== HF API CONFIG =================

HF_TOKEN = config("HF_TOKEN")

PRON_API_URL = "https://api-inference.huggingface.co/models/facebook/wav2vec2-base-960h"

headers = {
    "Authorization": f"Bearer {HF_TOKEN}"
}


# =============== PRONUNCIATION CHECK =================
sentences = [
    "The quick brown fox jumps over the lazy dog",
    "Hello, how are you doing today?",
    "Artificial intelligence is changing education",
    "This platform helps students improve skills",
    "I would like to book a table for two",
    "She sells sea shells by the sea shore",
    "What time does the train leave tomorrow?",
    "Learning English requires daily practice",
    "Technology and innovation shape the future",
    "Please upload your pronunciation recording"
]


def pronunciation_view(request):
    context = {"sentences": sentences}

    if request.method == "POST":
        expected_text = request.POST.get("expected_text")
        audio_file = request.FILES.get("audio_file")

        context["expected_text"] = expected_text

        if not audio_file:
            context["error"] = "Audio fayl yuklanmadi."
            return render(request, "ai_module/pronunciation_page.html", context)


        if audio_file.size > 1 * 1024 * 1024:
            context["error"] = "Audio hajmi 1 MB dan oshmasligi kerak."
            return render(request, "ai_module/pronunciation_page.html", context)

        try:
            # Faylni vaqtinchalik saqlash
            temp_file_path = default_storage.save(f"temp_audio/{audio_file.name}", audio_file)
            full_path = os.path.join(settings.MEDIA_ROOT, temp_file_path)

            # Matnni aniqlash
            recognizer = sr.Recognizer()
            with sr.AudioFile(full_path) as source:
                audio = recognizer.record(source)
                try:
                    result_text = recognizer.recognize_google(audio)
                except sr.UnknownValueError:
                    result_text = ""
                    context["error"] = "Gap tushunilmadi. Iltimos, yana urinib ko‘ring."
                except sr.RequestError:
                    result_text = ""
                    context["error"] = "Google API bilan bog‘lanishda xatolik yuz berdi."

            context["result_text"] = result_text

            # Matnlar o‘xshashligini hisoblash
            if result_text:
                expected_words = expected_text.lower().split()
                actual_words = result_text.lower().split()

                similarity = difflib.SequenceMatcher(None, expected_text.lower(), result_text.lower()).ratio()
                similarity_percent = round(similarity * 100, 2)
                wrong_words = [word for word in expected_words if word not in actual_words]

                context["similarity"] = similarity_percent
                context["wrong_words"] = wrong_words

            # Foydalanuvchining audiosini ko‘rsatish uchun
            context["user_audio"] = default_storage.open(temp_file_path)

        except Exception as e:
            context["error"] = f"Xatolik: {str(e)}"

    return render(request, "ai_module/pronunciation_page.html", context)



# =============== TEXT ANALYSIS =================

def translate_explanation(text):
    return text

def check_grammar_languagetool(text):
    url = "https://api.languagetoolplus.com/v2/check"
    payload = {
        'text': text,
        'language': 'en-US'
    }

    try:
        response = requests.post(url, data=payload, timeout=10)
    except Exception as e:
        return {
            "corrected_text": text,
            "message": f"Serverga ulanishda xato: {str(e)}",
            "corrections": []
        }

    if response.status_code == 200:
        result_json = response.json()
        corrections = []
        corrected_text = text

        if not result_json.get("matches"):
            return {
                "corrected_text": text,
                "message": "✅ Ofarin! Siz matnni to‘g‘ri yozdingiz.",
                "corrections": []
            }

        for match in result_json["matches"]:
            offset = match["offset"]
            length = match["length"]
            original = text[offset:offset + length]
            replacements = [r["value"] for r in match.get("replacements", [])]
            message_en = match["message"]
            message_uz = translate_explanation(message_en)

            corrections.append({
                "original": original,
                "replacements": replacements,
                "message_uz": message_uz
            })

            if replacements:
                replacement = replacements[0]
                corrected_text = corrected_text[:offset] + replacement + corrected_text[offset + length:]

        return {
            "corrected_text": corrected_text,
            "message": "Xatolar topildi. Quyidagilarni tuzatish tavsiya qilinadi:",
            "corrections": corrections
        }
    else:
        return {
            "corrected_text": text,
            "message": f"Xatolik: Status {response.status_code}. Response: {response.text}",
            "corrections": []
        }


def text_analysis(request):
    form = TextAnalysisForm()
    result = None
    error = None

    if request.method == 'POST':
        form = TextAnalysisForm(request.POST)
        if form.is_valid():
            input_text = form.cleaned_data['text']
            try:
                result = check_grammar_languagetool(input_text)
            except Exception as e:
                error = str(e)

    return render(request, 'ai_module/text_analysis_page.html', {
        'form': form,
        'result': result,
        'error': error
    })


# =============== INTERACTIVE TESTS =================

def interactive_tests(request):
    return render(request, 'ai_module/interactive_tests.html')


def get_quiz(request):
    level = request.GET.get('level', 'A1')

    file_path = os.path.join(
        os.path.dirname(__file__),
        'templates',
        'ai_module',
        'grammar_quiz.json'
    )

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        return JsonResponse({'error': f"JSON o‘qishda xato: {str(e)}"})

    all_quizzes = data.get(level, [])

    if len(all_quizzes) >= 5:
        selected = random.sample(all_quizzes, 5)
    else:
        selected = all_quizzes

    for q in selected:
        options = q["incorrect_answers"] + [q["correct_answer"]]
        random.shuffle(options)
        q["options"] = options

    return JsonResponse({'quiz': selected})

# =============== FLASHCARDS =================

def flashcards(request):
    return render(request, 'ai_module/flashcards.html')


def get_flashcards(request):
    level = request.GET.get('level', 'A1')

    file_path = os.path.join(
        os.path.dirname(__file__),
        'templates',
        'ai_module',
        'flashcards.json'
    )

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        return JsonResponse({'error': f"JSON o‘qishda xato: {str(e)}"})

    all_cards = data.get(level, [])

    if len(all_cards) >= 5:
        selected = random.sample(all_cards, 5)
    else:
        selected = all_cards

    all_translations = [c["translation"] for cards in data.values() for c in cards]

    for card in selected:
        correct = card["translation"]
        incorrects = random.sample(
            [t for t in all_translations if t != correct],
            k=min(3, len(all_translations) - 1)
        )
        options = incorrects + [correct]
        random.shuffle(options)
        card["options"] = options

    return JsonResponse({'cards': selected})


# =============== DIALOG EXERCISES =================

def get_dialog_questions(request):
    level = request.GET.get('level', 'A1')

    file_path = os.path.join(
        os.path.dirname(__file__),
        'templates',
        'ai_module',
        'dialog_questions.json'
    )

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        return JsonResponse({'error': f'JSON o‘qishda xato: {str(e)}'})

    questions = data.get(level, [])
    selected = random.sample(questions, min(5, len(questions)))

    return JsonResponse({'questions': selected})


def dialog_exercises_view(request):
    return render(request, 'ai_module/dialog_exercises.html')


@csrf_exempt
def check_dialog_answer(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            user_text = data.get("answer", "")
            result = check_grammar_languagetool(user_text)
            return JsonResponse(result)
        except Exception as e:
            return JsonResponse({"error": f"Server xatosi: {str(e)}"}, status=500)
    else:
        return JsonResponse({"error": "Faqat POST ishlaydi."})

# =============== VIDEO SUBTITLE VIEW =================

@csrf_exempt
def video_subtitle_view(request):
    videos = {
        'A1': [
            "https://www.youtube.com/embed/aNYEtGxjGVc",
            "https://www.youtube.com/embed/WyUZFcwRFdS",
            "https://www.youtube.com/embed/aSsi-kNzDXk",
        ],
        'A2': [
            "https://www.youtube.com/embed/G1PnQBEImqU",
            "https://www.youtube.com/embed/USXs-T8HpYU",
            "https://www.youtube.com/embed/fI62h9YwrMQ",
        ],
        'B1': [
            "https://www.youtube.com/embed/jQCxy43FiOQ",
            "https://www.youtube.com/embed/OA0gL2NrHuo",
            "https://www.youtube.com/embed/XdobJv1Kr9o",
        ],
        'B2': [
            "https://www.youtube.com/embed/0-__dGloMyw",
            "https://www.youtube.com/embed/Q2GPBoPUbtk",
            "https://www.youtube.com/embed/4bgC0w-LD9E",
        ],
        'C1': [
            "https://www.youtube.com/embed/8GlGrsO0cQE",
            "https://www.youtube.com/embed/3XxKg3cG8lE",
            "https://www.youtube.com/embed/4T_rB8rXnu4",
        ]
    }

    level = request.GET.get('level', 'A1')
    video_list = videos.get(level, [])
    video_url = random.choice(video_list) if video_list else None

    result = None
    subtitle = ""

    if request.method == "POST":
        level = request.POST.get('level', 'A1')
        subtitle = request.POST.get('subtitle', '')

        lang_tool_url = "https://api.languagetool.org/v2/check"
        payload = {
            'text': subtitle,
            'language': 'en-US'
        }
        r = requests.post(lang_tool_url, data=payload)

        corrections = []
        if r.status_code == 200:
            matches = r.json().get("matches", [])
            for match in matches:
                corrections.append({
                    'original': subtitle[match['offset']: match['offset'] + match['length']],
                    'replacements': [rep['value'] for rep in match.get('replacements', [])],
                    'message': match['message'],
                })

            if corrections:
                result = {
                    'message': "Xatolar topildi!",
                    'corrections': corrections
                }
            else:
                result = {
                    'message': "Zo'r! Subtitiringiz to'g'ri.",
                    'corrections': []
                }
        else:
            result = {
                'message': f"API xatosi: {r.status_code}",
                'corrections': []
            }

        video_list = videos.get(level, [])
        video_url = random.choice(video_list) if video_list else None

    context = {
        'level': level,
        'video_url': video_url,
        'result': result,
        'subtitle': subtitle,
    }
    return render(request, 'ai_module/video_subtitle.html', context)


# =============== AUDIO TEST VIEW =================

@csrf_exempt
def audio_test_view(request):
    result_text = None
    error = None

    if request.method == "POST":
        audio_file = request.FILES.get("audio_file")

        if audio_file:
            file_ext = os.path.splitext(audio_file.name)[1]
            temp_path = f"temp_audio{file_ext}"

            with open(temp_path, "wb") as f:
                for chunk in audio_file.chunks():
                    f.write(chunk)

            # MP3 -> WAV konvertatsiya
            if temp_path.endswith(".mp3"):
                sound = AudioSegment.from_mp3(temp_path)
                wav_path = temp_path.replace(".mp3", ".wav")
                sound.export(wav_path, format="wav")
                temp_path = wav_path

            try:
                recognizer = sr.Recognizer()
                with sr.AudioFile(temp_path) as source:
                    audio_data = recognizer.record(source)
                    result = recognizer.recognize_google(audio_data, language="en-US")
                    result_text = result

            except sr.UnknownValueError:
                error = "Nutqni aniqlab bo‘lmadi."
            except sr.RequestError as e:
                error = f"Google servisi bilan bog‘lanishda xatolik: {str(e)}"
            except Exception as e:
                error = str(e)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

    return render(request, "ai_module/audio_test.html", {
        "result_text": result_text,
        "error": error
    })

# =============== LISTENING FILL VIEW =================

@csrf_exempt
def listening_fill_view(request):
    import difflib

    user_text = ""
    result = None
    error = None

    levels = ["A1", "A2", "B1", "B2", "C1"]
    level = request.GET.get("level", "A1")

    json_path = os.path.join(
        os.path.dirname(__file__),
        'templates',
        'ai_module',
        'audio_transcripts.json'
    )

    audio_data = {}
    audio_list = []
    selected_audio = None

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            audio_data = json.load(f)
    except Exception as e:
        error = f"Audio JSON o‘qishda xato: {str(e)}"
        audio_data = {}

    if request.method == "POST":
        level = request.POST.get("level", "A1")
        audio_list = audio_data.get(level, [])

        selected_audio_index = request.POST.get("audio_index")
        if selected_audio_index is not None and audio_list:
            try:
                selected_audio = audio_list[int(selected_audio_index)]
            except:
                selected_audio = None
        elif audio_list:
            selected_audio = random.choice(audio_list)

        user_text = request.POST.get("user_text", "")

        # check only if user clicked the check button
        if request.POST.get("check_answer") == "1" and selected_audio:
            transcript_text = selected_audio["transcript"]

            def compare_transcripts(user_text, transcript_text):
                d = difflib.Differ()
                diff = list(d.compare(transcript_text.split(), user_text.split()))
                mistakes = []
                for item in diff:
                    if item.startswith('- '):
                        mistakes.append({
                            'type': 'missing',
                            'word': item[2:]
                        })
                    elif item.startswith('+ '):
                        mistakes.append({
                            'type': 'extra',
                            'word': item[2:]
                        })
                return mistakes

            transcript_diffs = compare_transcripts(user_text, transcript_text)
            grammar_result = check_grammar_languagetool(user_text)

            result = {
                "message": "Xatolar topildi!" if transcript_diffs or grammar_result.get("corrections") else "A'lo! Matn transcript bilan mos tushdi va xatosiz.",
                "diffs": transcript_diffs,
                "grammar": grammar_result.get("corrections", [])
            }
    else:
        audio_list = audio_data.get(level, [])
        if audio_list:
            selected_audio = random.choice(audio_list)

    context = {
        "levels": levels,
        "level": level,
        "audio_list": list(enumerate(audio_list)),
        "selected_audio": selected_audio,
        "user_text": user_text,
        "result": result,
        "error": error
    }

    return render(request, "ai_module/listening_fill.html", context)

# =============== GRAMATIK TEKSHIRUV =================

def extract_text_from_docx(file):
    text = ""
    doc = Document(file)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_pdf(file):
    text = ""
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    for page in pdf:
        text += page.get_text()
    return text

def grammar_check_start(request):
    matches = []
    user_text = ""

    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        if uploaded_file:
            filename = uploaded_file.name.lower()
            if filename.endswith('.txt'):
                user_text = uploaded_file.read().decode('utf-8')
            elif filename.endswith('.docx'):
                user_text = extract_text_from_docx(uploaded_file)
            elif filename.endswith('.pdf'):
                user_text = extract_text_from_pdf(uploaded_file)
            else:
                user_text = "Yuklangan fayl formati qo‘llab-quvvatlanmaydi."
        else:
            user_text = request.POST.get('input_text')

        if user_text.strip():
            url = "https://api.languagetool.org/v2/check"
            data = {
                'text': user_text,
                'language': 'en-US'
            }

            response = requests.post(url, data=data)
            result = response.json()

            for match in result.get('matches', []):
                message = match.get('message')
                offset = match.get('offset')
                length = match.get('length')
                context = user_text[offset:offset+length]
                replacements = [rep['value'] for rep in match.get('replacements', [])]
                matches.append({
                    'message': message,
                    'context': context,
                    'replacements': replacements
                })

    return render(request, 'ai_module/grammar_check.html', {
        'matches': matches,
        'user_text': user_text
    })

# =============== KONTEKSTUAL TARJIMA =================

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as pdf_doc:
        for page in pdf_doc:
            text += page.get_text()
    return text


def context_translate_start(request):
    translated_text = ""
    user_text = ""
    error_message = ""

    if request.method == 'POST':
        lang_pair = request.POST.get('language_pair', 'en|uz')

        # 1) text area matn bo'lsa
        user_text = request.POST.get('input_text', '').strip()

        # 2) yoki fayl yuklangan bo'lsa
        uploaded_file = request.FILES.get('uploaded_file')

        if uploaded_file:
            # Save file temporarily
            file_path = default_storage.save(uploaded_file.name, uploaded_file)

            file_ext = os.path.splitext(uploaded_file.name)[1].lower()

            if file_ext == '.docx':
                user_text = extract_text_from_docx(file_path)
            elif file_ext == '.pdf':
                user_text = extract_text_from_pdf(file_path)
            else:
                error_message = "Faqat PDF yoki DOCX formatidagi faylni yuklang."

            # remove temporary file
            if default_storage.exists(file_path):
                default_storage.delete(file_path)

        if user_text and not error_message:
            from_lang, to_lang = lang_pair.split('|')

            url = "https://api.mymemory.translated.net/get"
            params = {
                'q': user_text,
                'langpair': f'{from_lang}|{to_lang}'
            }

            try:
                response = requests.get(url, params=params, timeout=10)
                data = response.json()
                translated_text = data.get('responseData', {}).get('translatedText', '')

                if not translated_text:
                    error_message = "Tarjima natijasi olinmadi. Keyinroq urinib ko‘ring."

            except Exception as e:
                translated_text = ""
                error_message = "Tarjima jarayonida xatolik yuz berdi. Iltimos, keyinroq urinib ko‘ring."

    return render(request, 'ai_module/context_translate.html', {
        'user_text': user_text,
        'translated_text': translated_text,
        'error_message': error_message
    })

# =============== DARS REJASI =================
import json
from django.shortcuts import render
import os

def lesson_plans_view(request):
    json_path = os.path.join(
        os.path.dirname(__file__),
        "templates",
        "ai_module",
        "lesson_plans.json"
    )
    with open(json_path, "r", encoding="utf-8") as f:
        lesson_plans = json.load(f)

    # Foydalanuvchi tanlagan plan nomini olish
    selected_plan_name = request.GET.get("plan")

    selected_plan = None

    # Agar tanlangan bo'lsa, lesson_plans dan qidir
    if selected_plan_name and selected_plan_name in lesson_plans:
        selected_plan = lesson_plans[selected_plan_name]

    context = {
        "lesson_plans": lesson_plans,
        "selected_plan_name": selected_plan_name,
        "selected_plan": selected_plan,
    }
    return render(request, "ai_module/lesson_plan_generation.html", context)

# =============== LESSON RESOURCES =================
import os
import json
from django.shortcuts import render

def lesson_resources_view(request):
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    json_path = os.path.join(BASE_DIR, 'ai_module', 'templates', 'ai_module', 'lesson_resources.json')

    with open(json_path, encoding='utf-8') as f:
        data = json.load(f)

    topics = list(data.keys())
    selected_topic = None
    sites = None

    if request.method == "POST":
        selected_topic = request.POST.get("selected_topic")
        if selected_topic in data:
            sites = data[selected_topic]

    return render(request, "ai_module/lesson_resources.html", {
        "topics": topics,
        "selected_topic": selected_topic,
        "sites": sites
    })

import math
from django.shortcuts import render
from .forms import RadioactiveDecayForm


def radioactive_decay(request):
    form = RadioactiveDecayForm(request.POST or None)

    result = None
    if request.method == "POST" and form.is_valid():
        N0 = form.cleaned_data["N0"]
        lambd = form.cleaned_data["lambd"]
        t_max = form.cleaned_data["t_max"]
        steps = form.cleaned_data["steps"]

        # vaqt nuqtalari
        dt = t_max / (steps - 1)
        t_values = [round(i * dt, 6) for i in range(steps)]

        # yechim: N(t) = N0 * exp(-λ t)
        N_values = [N0 * math.exp(-lambd * t) for t in t_values]

        # yarim parchalanish vaqti
        half_life = math.log(2) / lambd

        # jadval uchun (yuqori aniqlikni ko‘rsatish)
        table = [{"t": t_values[i], "N": round(N_values[i], 6)} for i in range(steps)]

        result = {
            "half_life": round(half_life, 6),
            "t_values": t_values,
            "N_values": [round(x, 6) for x in N_values],
            "table": table,
        }

    return render(request, "ai_module/radioactive_decay.html", {"form": form, "result": result})